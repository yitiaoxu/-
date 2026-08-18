# -*- coding: utf-8 -*-
"""Generate resume project summary Word document."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

OUTPUT = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    "resume_project_plate_vision.docx",
)

TITLE = "\u57fa\u4e8e RK3568 \u7684 FPGA+PCIe \u8f66\u724c\u68c0\u6d4b\u8bc6\u522b\u5d4c\u5165\u5f0f\u4e0a\u4f4d\u673a\u7cfb\u7edf"
SUBTITLE = "\u7b80\u5386\u9879\u76ee\u603b\u7ed3 \u00b7 \u53ef\u76f4\u63a5\u590d\u5236\u6216\u6309\u9700\u4fee\u6539"


def set_doc_font(doc):
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_title(doc, text, size=18):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)


def add_para(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.35


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.name = "Microsoft YaHei"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.35


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = val
    doc.add_paragraph()


def main():
    doc = Document()
    set_doc_font(doc)
    for sec in doc.sections:
        sec.top_margin = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(2.8)
        sec.right_margin = Cm(2.8)

    add_title(doc, TITLE)
    add_subtitle(doc, SUBTITLE)
    doc.add_paragraph()

    add_heading(doc, "\u4e00\u3001\u9879\u76ee\u6982\u8ff0")
    add_para(
        doc,
        "\u9762\u5411\u4ea4\u901a/\u5b89\u9632\u573a\u666f\u7684\u5d4c\u5165\u5f0f\u8f66\u724c\u8bc6\u522b\u7cfb\u7edf\u3002"
        "FPGA \u91c7\u96c6 1280\u00d7720 \u89c6\u9891\u7ecf PCIe DMA \u9001\u5165 RK3568 \u8fb9\u7f18\u7aef\uff0c"
        "\u91c7\u7528 RKNN \u5728 NPU \u4e0a\u8fd0\u884c\u68c0\u6d4b+\u8bc6\u522b\u53cc\u6a21\u578b\uff1b"
        "Qt5 \u4e0a\u4f4d\u673a\u5b8c\u6210\u5b9e\u65f6\u9884\u89c8\u3001\u8f66\u724c\u53f7\u5c55\u793a\u3001"
        "\u6279\u91cf\u8bc4\u6d4b\u4e0e\u5168\u94fe\u8def\u5ef6\u8fdf\u5206\u6790\u3002",
    )
    add_para(
        doc,
        "\u4e00\u53e5\u8bdd\u63cf\u8ff0\uff1a\u5728 RK3568 \u8fb9\u7f18\u7aef\u5b9e\u73b0 FPGA \u7ecf PCIe "
        "\u5b9e\u65f6\u9001\u56fe\u3001NPU \u4e24\u9636\u6bb5\u8f66\u724c\u68c0\u6d4b\u4e0e\u8bc6\u522b\uff0c"
        "\u5e76\u914d\u5957 Qt \u5de5\u63a7\u4e0a\u4f4d\u673a\uff0c\u5355\u5e27\u7b97\u6cd5\u5ef6\u8fdf\u7ea6 60\u201380 ms\u3002",
    )

    add_heading(doc, "\u4e8c\u3001\u4e3b\u8981\u804c\u8d23\u4e0e\u9879\u76ee\u6210\u679c")
    add_bullets(
        doc,
        [
            "\u8bbe\u8ba1\u5e76\u5b9e\u73b0\u68c0\u6d4b+\u8bc6\u522b\u4e24\u9636\u6bb5 RKNN \u63a8\u7406 Pipeline\uff08C++/OpenCV\uff09\uff0c"
            "\u652f\u6301 INT8/FP \u6a21\u578b\u81ea\u52a8\u5207\u6362\u3001\u4e09\u5934 INT8 \u89e3\u7801\u3001NMS \u540e\u5904\u7406\u53ca"
            "\u4e2d\u6587\u8f66\u724c\u53f7/\u989c\u8272\u8f93\u51fa\uff1bQt \u4e0a\u4f4d\u673a\u4e0e CLI \u5171\u7528\u540c\u4e00\u5957\u6838\u5fc3\u7b97\u6cd5\u6a21\u5757\u3002",
            "\u5f00\u53d1 Qt5 \u5de5\u63a7\u98ce\u4e0a\u4f4d\u673a\uff08vision_qt_demo\uff09\uff0c\u652f\u6301 PCIe \u5b9e\u65f6\u91c7\u96c6\u3001"
            "\u5355\u56fe/\u6587\u4ef6\u5939/\u89c6\u9891\u6279\u91cf\u5904\u7406\u3001\u5de6\u4fa7\u5b9e\u65f6\u8f66\u724c\u53f7\u5c55\u793a\u3001"
            "process_wall / \u7aef\u5230\u7aef\u5ef6\u8fdf\u663e\u793a\u53ca\u7ed3\u679c\u753b\u5eca\u7ffb\u9875\u6d4f\u89c8\u3002",
            "\u5b9e\u73b0 PCIe \u89c6\u9891\u91c7\u96c6\u94fe\u8def\uff1aLinux \u5b57\u7b26\u8bbe\u5907\u9a71\u52a8\uff08pango_pci_driver\uff09+ "
            "\u7528\u6237\u6001 DMA \u8bfb\u5e27\uff08\u6574\u5e27/\u9010\u884c mmap\u3001RGB565/RGB888 \u8f6c BGR\uff09\uff0c"
            "\u4e0e FPGA \u4fa7 1280\u00d7720 \u89c6\u9891\u6d41\u5bf9\u63a5\u3002",
            "\u4f18\u5316\u5b9e\u65f6\u6027\u4e0e\u7a33\u5b9a\u6027\uff1a\u5e27\u663e\u793a\u5408\u5e76\u5237\u65b0\u907f\u514d UI \u963b\u585e\uff1b"
            "\u8f66\u724c\u7ed3\u679c\u53d8\u5316\u9a71\u52a8\u754c\u9762\u66f4\u65b0\uff1b\u68c0\u6d4b\u6846 EMA \u65f6\u5e8f\u5e73\u6ed1\uff1b"
            "\u5efa\u7acb\u5168\u94fe\u8def\u5206\u6bb5\u5ef6\u8fdf Profiling\uff08DMA/\u8f6c\u8272/letterbox/det NPU/"
            "rec NPU/draw/QImage/e2e\uff09\u3002",
            "\u642d\u5efa\u79bb\u7ebf\u9a8c\u8bc1\u5de5\u5177\u94fe\uff08rknn_infer_one + \u6279\u91cf\u811a\u672c\uff09\uff0c"
            "\u652f\u6301\u968f\u673a\u62bd\u6837\u3001CSV \u6c47\u603b\u3001\u5206\u9636\u6bb5\u8017\u65f6\u7edf\u8ba1\u3002",
            "\u6027\u80fd\u6307\u6807\uff08RK3568 \u5b9e\u6d4b\uff09\uff1a\u68c0\u6d4b\u7ea6 52 ms\u3001\u8bc6\u522b\u7ea6 3 ms\u3001"
            "\u753b\u6846\u7ea6 5 ms\uff0c\u7b97\u6cd5\u603b\u8017\u65f6\uff08process_wall\uff09\u7ea6 60\u201370 ms\uff0c"
            "\u542b PCIe \u53d6\u56fe\u4e0e\u8f6c\u8272\u7684\u7aef\u5230\u7aef\u7ea6 80 ms\u3002",
        ],
    )

    add_heading(doc, "\u4e09\u3001\u6280\u672f\u6808")
    add_table(
        doc,
        ["\u7c7b\u522b", "\u6280\u672f"],
        [
            ["\u8bed\u8a00/\u6846\u67b6", "C++14\u3001Qt5\u3001OpenCV"],
            ["AI \u63a8\u7406", "RKNN\u3001RK3568 NPU\u3001INT8 \u91cf\u5316\u3001YOLO \u7c7b\u68c0\u6d4b + \u5e8f\u5217\u8bc6\u522b"],
            ["\u5d4c\u5165\u5f0f/\u9a71\u52a8", "Linux \u5185\u6838\u6a21\u5757\u3001PCIe DMA\u3001ioctl\u3001mmap\u3001ARM64"],
            ["\u786c\u4ef6\u534f\u540c", "FPGA \u89c6\u9891\u91c7\u96c6\u3001RGB565/RGB888\u3001ROI \u63a5\u53e3"],
            ["\u5de5\u7a0b\u5316", "CMake\u3001Shell \u90e8\u7f72\u811a\u672c\u3001CLI \u6279\u91cf\u8bc4\u6d4b\u3001CSV \u6307\u6807\u5bfc\u51fa"],
        ],
    )

    add_heading(doc, "\u56db\u3001\u7cfb\u7edf\u67b6\u6784\uff08\u7b80\u8ff0\uff09")
    add_bullets(
        doc,
        [
            "\u89c6\u9891\u6e90\uff1aFPGA HDMI 1280\u00d7720 \u2192 PCIe DMA \u2192 RK3568 \u7528\u6237\u6001",
            "\u63a8\u7406\uff1aletterbox(640) \u2192 plate_detect(RKNN) \u2192 plate_rec(RKNN) \u2192 \u753b\u6846/\u4e2d\u6587\u6807\u7b7e",
            "\u5c55\u793a\uff1aQt \u4e3b\u7a97\u53e3\u5b9e\u65f6\u9884\u89c8 + \u5de6\u4fa7\u8f66\u724c\u53f7\u9762\u677f + \u9876\u680f\u5ef6\u8fdf\u6307\u6807",
            "\u90e8\u7f72\uff1arebuild.sh / run.sh \u4e00\u952e\u7f16\u8bd1\u8fd0\u884c\uff0cweights/ \u7ba1\u7406 RKNN \u6a21\u578b",
        ],
    )

    add_heading(doc, "\u4e94\u3001\u7cbe\u7b80\u7248\uff08\u7b80\u5386\u7bc7\u5e45\u7d27\u5f20\u65f6\u4f7f\u7528\uff09")
    add_bullets(
        doc,
        [
            "\u57fa\u4e8e RK3568 + RKNN NPU\uff0c\u5b9e\u73b0\u8f66\u724c\u68c0\u6d4b/\u8bc6\u522b\u4e24\u9636\u6bb5\u63a8\u7406 Pipeline\uff0c"
            "\u5355\u5e27\u7b97\u6cd5\u5ef6\u8fdf\u7ea6 60 ms\u3002",
            "\u5f00\u53d1 Qt5 \u5d4c\u5165\u5f0f\u4e0a\u4f4d\u673a\uff0c\u96c6\u6210 FPGA PCIe DMA 1280\u00d7720 \u5b9e\u65f6\u89c6\u9891\u6d41\u3001"
            "\u6279\u91cf\u56fe\u7247/\u89c6\u9891\u5904\u7406\u3001\u4e2d\u6587\u8f66\u724c\u53ef\u89c6\u5316\u4e0e\u5ef6\u8fdf\u76d1\u63a7\u3002",
            "\u7f16\u5199 Linux PCIe \u9a71\u52a8\u4e0e\u7528\u6237\u6001\u91c7\u96c6\u6a21\u5757\uff0c\u5e76\u5b8c\u6210 UI \u6027\u80fd\u4f18\u5316\u3001"
            "\u68c0\u6d4b\u6846\u65f6\u5e8f\u5e73\u6ed1\u53ca\u5168\u94fe\u8def\u5ef6\u8fdf Profiling\u3002",
        ],
    )

    add_heading(doc, "\u516d\u3001\u586b\u5199\u5efa\u8bae")
    add_bullets(
        doc,
        [
            "\u8865\u5145\u9879\u76ee\u8d77\u6b62\u65f6\u95f4\uff08\u5982 2025.03 \u2013 2026.07\uff09\u53ca\u4e2a\u4eba\u89d2\u8272\uff08\u72ec\u7acb/\u56e2\u961f\uff09\u3002",
            "\u6027\u80fd\u6570\u636e\u8bf7\u4ee5\u677f\u7aef\u5b9e\u6d4b\u4e3a\u51c6\uff0c\u9762\u8bd5\u65f6\u53ef\u7ed3\u5408 [latency] \u5206\u6bb5\u65e5\u5fd7\u8bf4\u660e\u4f18\u5316\u65b9\u5411\u3002",
            "\u82e5\u9a71\u52a8/FPGA \u4e3a\u534f\u4f5c\u5b8c\u6210\uff0c\u5efa\u8bae\u8868\u8ff0\u4e3a\u300c\u53c2\u4e0e PCIe \u91c7\u96c6\u94fe\u8def\u8054\u8c03\u300d\u3002",
        ],
    )

    doc.save(OUTPUT)
    print("Saved:", OUTPUT)


if __name__ == "__main__":
    main()
