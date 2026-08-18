# -*- coding: utf-8 -*-
"""Generate latency / optimization / INT8 report Word document."""

import os
import sys

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, SCRIPT_DIR)
import report_content as C

OUTPUT = os.path.join(os.path.dirname(SCRIPT_DIR), "latency_optimization_int8_report.docx")


def set_font(run, size=11, bold=False, color=None):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run(text), 20, True)


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run(text), 10, False, RGBColor(0x66, 0x66, 0x66))


def add_h1(doc, text):
    h = doc.add_heading(text, level=1)
    for r in h.runs:
        set_font(r, 14, True)


def add_h2(doc, text):
    h = doc.add_heading(text, level=2)
    for r in h.runs:
        set_font(r, 12, True)


def add_p(doc, text):
    p = doc.add_paragraph()
    set_font(p.add_run(text))
    p.paragraph_format.line_spacing = 1.35
    p.paragraph_format.space_after = Pt(6)


def add_bullets(doc, items):
    for t in items:
        p = doc.add_paragraph(t, style="List Bullet")
        for r in p.runs:
            set_font(r)
        p.paragraph_format.line_spacing = 1.3


def add_table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = str(h)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = str(val)
    doc.add_paragraph()


def add_code(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    p.paragraph_format.left_indent = Cm(0.5)


def render_block(doc, block):
    if block.get("h1"):
        add_h1(doc, block["h1"])
    if block.get("h2"):
        add_h2(doc, block["h2"])
    if block.get("table"):
        headers, rows = block["table"]
        add_table(doc, headers, rows)
    for para in block.get("paras", []):
        add_p(doc, para)
    if block.get("bullets"):
        add_bullets(doc, block["bullets"])
    if block.get("code"):
        add_code(doc, block["code"])
    for para in block.get("paras_after", []):
        add_p(doc, para)


def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    for sec in doc.sections:
        sec.top_margin = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(2.8)
        sec.right_margin = Cm(2.8)

    add_title(doc, C.TITLE)
    add_subtitle(doc, C.SUBTITLE)
    doc.add_paragraph()

    for block in C.SECTIONS:
        render_block(doc, block)

    add_p(doc, u"\u62a5\u544a\u6587\u4ef6\uff1alatency_optimization_int8_report.docx")
    doc.save(OUTPUT)
    print("Saved:", OUTPUT)


if __name__ == "__main__":
    main()
